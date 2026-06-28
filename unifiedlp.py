import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

# --- 1. GLOBAL LOCALIZATION METRICS ---
LANG_MAP = {
    "English": {
        "title_prefix": "LESSON PLAN",
        "resources_title": "RESOURCES & MATERIALS",
        "resources_body": "Smart board, Chromebook, Writing table, Projector, Screen share with laptop",
        "keywords_key": "KEYWORDS",
        "pedati_key": "PEDATI",
        "tbl_hdr": ['STAGE (PEDATI)', ' ACTIVITY 1 ', ' ACTIVITY 2 '],
        "hod_title": "HOD APPROVAL & REMARKS",
        "hod_headers": [["REMARK", "SIGNATURE / STAMP"], ["DATE:", "NAME:"]],
        "page_lbl": "Page "
    },
    "Malay": {
        "title_prefix": "RANCANGAN PENGAJARAN",
        "resources_title": "SUMBER & BAHAN BANTU MENGAJAR",
        "resources_body": "Papan pintar, Chromebook, Meja tulis, Projektor, Perkongsian skrin dengan komputer riba",
        "keywords_key": "KATA KUNCI",
        "pedati_key": "PEDATI",
        "tbl_hdr": ['PERINGKAT (PEDATI)', ' AKTIVITI 1 ', ' AKTIVITI 2 '],
        "hod_title": "PENGESAHAN & ULASAN KETUA JABATAN",
        "hod_headers": [["ULASAN", "TANDATANGAN / COP"], ["TARIKH:", "NAMA:"]],
        "page_lbl": "Halaman "
    }
}

st.set_page_config(page_title="UNIFIED Lesson Planners", layout="wide")
st.title("🎓 UNIFIED SMART LESSON PLANNER")

# --- SIDEBAR CONTROL BAR ---
st.sidebar.header("🛠️ PORTAL SETTINGS")

user_api_key = st.sidebar.text_input(
    "🔑 ENTER YOUR GEMINI API KEY:", 
    type="password", 
    help="Create and get your API key from Google AI Studio."
)

selected_platform = st.sidebar.selectbox(
    "📂 SELECT LESSON PLANNER TYPE:",
    ["PEDATI LP", "UNIVERSAL LP", "MERGED LP", "Secondary LP"]
)

selected_lang = st.sidebar.selectbox(
    "🌐 SELECT TARGET LANGUAGE:",
    ["English", "Malay"]
)

def get_working_model(api_key):
    try:
        genai.configure(api_key=api_key)
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                if "flash" in m.name:
                    return m.name
        return "models/gemini-1.5-flash"
    except Exception as e:
        st.sidebar.error(f"INVALID API KEY: {str(e)}")
        return None

selected_model_name = None
if user_api_key:
    selected_model_name = get_working_model(user_api_key)
    if selected_model_name:
        st.sidebar.success(f"CONNECTED: {selected_model_name.upper()}")
else:
    st.sidebar.warning("⚠️ ENTER YOUR OWN GEMINI API KEY BEFORE USING THIS PORTAL.")


# --- 2. MULTI-TEMPLATE AI GENERATION ENGINE ---
def generate_lesson_plan(topic, syllabus, extra_context, api_key, model_name, platform, lang):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    
    lang_instruction = "Generate the entire lesson plan in English. NO Malay terms." if lang == "English" else "Penulisan penuh adalah di dalam BAHASA MELAYU sahaja. JANGAN gunakan istilah Bahasa Inggeris."
    
    # Structural Prompt Setup
    if lang == "English":
        core_criteria = """
SECTION: LESSON OBJECTIVES
1. [Objective 1]
2. [Objective 2]
3. [Objective 3]
4. [Objective 4]

SECTION: LESSON OUTCOMES
1. [Outcome 1]
2. [Outcome 2]
3. [Outcome 3]
4. [Outcome 4]

SECTION: SUCCESS CRITERIA
1. [Criteria 1]
2. [Criteria 2]
3. [Criteria 3]
4. [Criteria 4]

SECTION: PREREQUISITE
1. [Prerequisite description]

SECTION: KEYWORDS
[Provide 6 standalone technical terms separated by commas or lines without numbers or bullets]

SECTION: HOTS
1. [HOTS point 1]
2. [HOTS point 2]
3. [HOTS point 3]
4. [HOTS point 4]"""
        
        dig_cit = """\n
SECTION: DIGITAL CITIZENSHIP
1. [Digital citizenship point 1]
2. [Digital citizenship point 2]
3. [Digital citizenship point 3]
4. [Digital citizenship point 4]"""
        
        pedati_stages = """\n
SECTION: PEDATI STAGES
STAGE: P [PRIOR KNOWLEDGE] | CB: [Activity] | SB: [Activity]
STAGE: E [ENGAGE] | CB: [Activity] | SB: [Activity]
STAGE: D [DEVELOP] | CB: [Activity] | SB: [Activity]
STAGE: A [APPLY] | CB: [Activity] | SB: [Activity]
STAGE: T [TEST] | CB: [Activity] | SB: [Activity]
STAGE: I [IMPROVE] | CB: [Activity] | SB: [Activity]"""
        
        universal_blocks = """\n
SECTION: OPENING LESSON CONTENT
1. [Hook activity and transition plan]

SECTION: DIFFERENTIATION STRATEGIES (GREEN)
1. HA (Higher Achiever): [1 challenging activity]

SECTION: DIFFERENTIATION STRATEGIES (YELLOW)
1. MA (Medium Achiever): [1 core activity]

SECTION: DIFFERENTIATION STRATEGIES (RED)
1. LA (Lower Achiever): [1 scaffolded activity]

SECTION: BLENDED LEARNING ACTIVITY ONE (15 MINS)
1. Activity 1: [Descriptions]
2. Teacher Preparation: [Step-by-step before lesson]
3. Objectives: [3 points]
4. Student Tasks: [Step-by-step details]

SECTION: BLENDED LEARNING ACTIVITY TWO (15 MINS)
1. Activity 2: [Descriptions]
2. Teacher Preparation: [Step-by-step before lesson]
3. Objectives: [3 points]
4. Student Tasks: [Step-by-step details]
    
SECTION: PLENARY (EXIT TICKET)
1. [2-3 minute closing activity]

SECTION: HOMEWORK
1. [Task assigned based on topic]

SECTION: SUGGESTED WAY FORWARD TASK
1. [Hook activity and transition plan for next day lesson]"""
        
    else: # MALAY TEMPLATE DESIGN ROUTING RULES
        core_criteria = """
SECTION: OBJEKTIF PEMBELAJARAN
1. [Objektif 1]
2. [Objektif 2]
3. [Objektif 3]
4. [Objektif 4]

SECTION: HASIL PEMBELAJARAN
1. [Hasil 1]
2. [Hasil 2]
3. [Hasil 3]
4. [Hasil 4]

SECTION: KRITERIA KEJAYAAN
1. [Kriteria 1]
2. [Kriteria 2]
3. [Kriteria 3]
4. [Kriteria 4]

SECTION: PENGETAHUAN SEDIA ADA
1. [Pengetahuan sedia ada]

SECTION: KATA KUNCI
[Sediakan 6 item kosa kata sahaja tanpa sebarang nombor siri atau simbol bullet]

SECTION: KBAT
1. [KBAT poin 1]
2. [KBAT poin 2]
3. [KBAT poin 3]
4. [KBAT poin 4]"""
        
        dig_cit = """\n
SECTION: KEWARGANEGARAAN DIGITAL
1. [Kewarganegaraan digital poin 1]
2. [Kewarganegaraan digital poin 2]
3. [Kewarganegaraan digital poin 3]
4. [Kewarganegaraan digital poin 4]"""
        
        pedati_stages = """\n
SECTION: PERINGKAT PEDATI
STAGE: P [PRIOR KNOWLEDGE] | CB: [Aktiviti] | SB: [Aktiviti]
STAGE: E [ENGAGE] | CB: [Aktiviti] | SB: [Aktiviti]
STAGE: D [DEVELOP] | CB: [Aktiviti] | SB: [Aktiviti]
STAGE: A [APPLY] | CB: [Aktiviti] | SB: [Aktiviti]
STAGE: T [TEST] | CB: [Aktiviti] | SB: [Aktiviti]
STAGE: I [IMPROVE] | CB: [Aktiviti] | SB: [Aktiviti]"""
        
        universal_blocks = """\n
SECTION: KANDUNGAN PEMBUKAAN PELAJARAN
1. [Aktiviti Set Induksi dan pelan peralihan]

SECTION: STRATEGI PERBEZAAN (HIJAU)
1. HA (Higher Achiever): [1 aktiviti mencabar]

SECTION: STRATEGI PERBEZAAN (KUNING)
1. MA (Medium Achiever): [1 aktiviti teras]

SECTION: STRATEGI PERBEZAAN (MERAH)
1. LA (Lower Achiever): [1 aktiviti sokongan/scaffolded]

SECTION: AKTIVITI PEMBELAJARAN TERADUN SATU (15 MINIT)
1. Aktiviti 1: [Penerangan]
2. Persediaan Guru: [Langkah demi langkah sebelum pelajaran]
3. Objektif: [3 mata/poin]
4. Tugasan Pelajar: [Butiran langkah demi langkah]

SECTION: AKTIVITI PEMBELAJARAN TERADUN DUA (15 MINIT)
1. Aktiviti 2: [Penerangan]
2. Persediaan Guru: [Langkah demi langkah sebelum pelajaran]
3. Objektif: [3 mata/poin]
4. Tugasan Pelajar: [Butiran langkah demi langkah]
    
SECTION: PLENARI (TIKET KELUAR)
1. [Aktiviti penutup 2-3 minit]

SECTION: KERJA RUMAH
1. [Tugasan diberikan berdasarkan topik]

SECTION: CADANGAN TUGASAN UTK KELAS AKAN DATANG
1. [Aktiviti ransangan dan pelan peralihan untuk pelajaran hari esok]"""

    # Assemble and fine-tune structural prompt guidelines
    prompt = f"Topic: {topic}. Syllabus Code: {syllabus}. Context: {extra_context}.\n{lang_instruction}\n\n"
    prompt += """CRITICAL FORMATTING RULES:
1. DO NOT use double asterisks (**) anywhere.
2. DO NOT use bullet points (-) for listings. You must use sequential numbered listings (1., 2., 3., 4.) for all content segments.
3. EXCEPTION FOR KEYWORDS / KATA KUNCI: Do NOT use any numbers, list markers, or bullets. Just list the words.
4. ABSOLUTE MALAY LANGUAGE COMPLIANCE: Jangan gunakan perkataan 'MURID'. Menyeluruh kontek dalam Bahasa Melayu sahaja, digantikan dengan perkataan 'PELAJAR'.
5. Every section marker MUST start explicitly on a new line with 'SECTION: ' followed by the uppercase title.
6. STRIKT KEPATUHAN STRUKTUR: You are strictly FORBIDDEN from creating custom essay headings (like PENGENALAN, CIRI-CIRI, CABARAN, KESIMPULAN). You MUST ONLY generate content for the exact SECTION blocks provided below. Do not omit any blocks!\n\n"""

    if platform == "PEDATI Plan":
        prompt += core_criteria + dig_cit + pedati_stages
    elif platform == "UNIVERSAL Plan":
        prompt += core_criteria + dig_cit + universal_blocks
    elif platform == "MERGED Plan":
        prompt += core_criteria + dig_cit + universal_blocks + pedati_stages
    elif platform == "UNIVERSAL (No Dig-Cit)":
        prompt += core_criteria + universal_blocks

    try:
        response = model.generate_content(prompt)
        if response.candidates and response.candidates[0].content.parts:
            # Universal string cleaning for absolute compliance
            clean_text = response.text.replace("**", "")
            clean_text = clean_text.replace("Murid", "Pelajar").replace("murid", "pelajar").replace("MURID", "PELAJAR")
            return clean_text
        else:
            return "⚠️ The AI returned an empty response. Please wait 60 seconds and try again."
    except Exception as e:
        return f"System Error: {str(e)}"


def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


# --- 3. DYNAMIC WORD DOCUMENT EXPORT ENGINE ---
def create_word_export(topic, syllabus, text, lang):
    meta = LANG_MAP[lang]
    doc = Document()
    
    # Enforce Letter Geometry Profile & 0.5" Margins
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11.5)
    section.top_margin = section.bottom_margin = Inches(0.5)
    section.left_margin = section.right_margin = Inches(0.5)
    
    # Page Tracking Setup
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = 2
    f_run = footer_p.add_run(meta["page_lbl"])
    f_run.font.name, f_run.font.size = 'Arial Narrow', Pt(10)
    add_page_number(f_run)

    # Core Typography Defaults Override (12pt Base)
    style = doc.styles['Normal']
    style.font.name, style.font.size = 'Arial Narrow', Pt(12)
    p_format = style.paragraph_format
    p_format.line_spacing, p_format.space_after, p_format.space_before = 1.0, Pt(0), Pt(0)

    # Document Header Title - 14PT BOLD CAPITALS
    title_p = doc.add_paragraph()
    run_title = title_p.add_run(f'{meta["title_prefix"]}: {topic.upper()} ({syllabus.upper()})')
    run_title.bold = True
    run_title.font.size = Pt(14)
    title_p.paragraph_format.space_after = Pt(6)

    # Admin Table Setup
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["WEEK NO:", "DATE:"], ["NO. OF STUDENTS:", "DAY:"], ["VENUE / LAB NO:", "DURATION (MINS):"]] if lang == "English" else [["MINGGU:", "TARIKH:"], ["BIL. PELAJAR:", "HARI:"], ["TEMPAT / NO MAKMAL:", "TEMPOH (MINIT):"]]
    for r in range(3):
        for col_idx, text_lbl in [(0, labels[r][0]), (2, labels[r][1])]:
            cell_p = admin_table.cell(r, col_idx).paragraphs[0]
            r_run = cell_p.add_run(text_lbl)
            r_run.bold = True
            r_run.font.size = Pt(12)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Resources Setup
    p_res = doc.add_paragraph()
    run_res = p_res.add_run(meta["resources_title"])
    run_res.bold = True
    run_res.font.size = Pt(14)
    p_res.paragraph_format.space_after = Pt(4)
    
    res_table = doc.add_table(rows=1, cols=1)
    res_table.style = 'Table Grid'
    r_cell = res_table.cell(0, 0).paragraphs[0].add_run(meta["resources_body"])
    r_cell.font.size = Pt(12)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Content Processing Core Loop Engine Block
    sections = text.split('SECTION:')
    for section in sections:
        if not section.strip(): 
            continue
            
        lines = section.strip().split('\n')
        title = lines[0].strip().upper().replace("**", "")
        content_lines = lines[1:]

        # Create Headings (Forced 14pt Bold Capitals)
        p_sec = doc.add_paragraph()
        run_sec_title = p_sec.add_run(title)
        run_sec_title.bold = True
        run_sec_title.font.size = Pt(14)
        p_sec.paragraph_format.space_after = Pt(4)

        # 1. Custom 3x2 Matrix Processing for Keywords / Kata Kunci
        if meta["keywords_key"] in title:
            keywords_list = []
            # Split by commas or lines to extract words cleanly
            raw_content = " ".join(content_lines).replace(",", " ")
            for item in raw_content.split():
                item = item.strip().replace(".", "").replace("-", "")
                if item and not item.isdigit():
                    keywords_list.append(item)
            
            while len(keywords_list) < 6: keywords_list.append("")
            keywords_list = keywords_list[:6]
            
            key_table = doc.add_table(rows=2, cols=3)
            key_table.style = 'Table Grid'
            kw_idx = 0
            for row_i in range(2):
                for col_i in range(3):
                    cell_p = key_table.cell(row_i, col_i).paragraphs[0]
                    cell_p.alignment = 1
                    k_run = cell_p.add_run(keywords_list[kw_idx])
                    k_run.font.size = Pt(12)
                    kw_idx += 1
            
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # 2. Custom 3-Column Table Grid processing for PEDATI stages
        elif "|" in section and meta["pedati_key"] in title:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            
            for idx, h_text in enumerate(meta["tbl_hdr"]):
                hrun = hdr[idx].paragraphs[0].add_run(h_text)
                hrun.bold = True
                hrun.font.size = Pt(12)

            for line in content_lines:
                cleaned_line = line.replace("**", "").strip()
                if "|" in cleaned_line:
                    p = cleaned_line.split("|")
                    if len(p) >= 3:
                        row_cells = table.add_row().cells
                        col_0 = p[0].split(":")[-1].strip().upper() if ":" in p[0] else p[0].strip().upper()
                        col_1 = p[1].split(":")[-1].strip() if ":" in p[1] else p[1].strip()
                        col_2 = p[2].split(":")[-1].strip() if ":" in p[2] else p[2].strip()
                        
                        for c_idx, c_text in enumerate([col_0, col_1, col_2]):
                            cell_p = row_cells[c_idx].paragraphs[0]
                            c_run = cell_p.add_run(c_text)
                            c_run.font.size = Pt(12)
                            if c_idx == 0: c_run.bold = True
            
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            
        # 3. Standard Grid Box Structures Default Block Configuration (Forced 12pt Content)
        else:
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell_p = table.cell(0, 0).paragraphs[0]
            cleaned_body = "\n".join([l.strip() for l in content_lines if l.strip()]).replace("**", "")
            body_run = cell_p.add_run(cleaned_body)
            body_run.font.size = Pt(12)
            
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 4. HOD Verification Layout
    doc.add_page_break()
    p_hod = doc.add_paragraph()
    run_hod = p_hod.add_run(meta["hod_title"])
    run_hod.bold = True
    run_hod.font.size = Pt(14)
    p_hod.paragraph_format.space_after = Pt(4)
    
    hod_table = doc.add_table(rows=3, cols=2)
    hod_table.style = 'Table Grid'
    
    headers_hod = [
        (meta["hod_headers"][0][0], 0, 0), (meta["hod_headers"][0][1], 0, 1),
        (meta["hod_headers"][1][0], 2, 0), (meta["hod_headers"][1][1], 2, 1)
    ]
    for text_val, r_i, c_i in headers_hod:
        hrun = hod_table.cell(r_i, c_i).paragraphs[0].add_run(text_val)
        hrun.bold = True
        hrun.font.size = Pt(12)
        
    hod_table.rows[1].height = Pt(40)

    # Tight paragraph formatting overrides across tables
    for t in [admin_table, hod_table]:
        for row in t.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.line_spacing = 1.0
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- 4. MAIN GUI INTERFACE ---
st.subheader(f"📝 ACTIVE PORTAL: {selected_platform.upper()} ({selected_lang.upper()})")
u_topic = st.text_input("LESSON TOPIC:")
u_syllabus = st.text_input("SYLLABUS CODE:")
u_extra = st.text_area("SPECIFIC CONTEXT / KEYWORDS (OPTIONAL):")

if st.button("🚀 GENERATE LESSON PLAN", type="primary"):
    if not user_api_key:
        st.error("❌ KEY CONFIGURATION ERROR! PLEASE INPUT YOUR GEMINI API KEY IN THE SIDEBAR.")
    elif not u_topic or not u_syllabus:
        st.error("❌ PLEASE PROVIDE BOTH A LESSON TOPIC AND A SYLLABUS CODE.")
    else:
        with st.spinner("AI IS BUILDING YOUR SELECTION..."):
            result = generate_lesson_plan(u_topic, u_syllabus, u_extra, user_api_key, selected_model_name, selected_platform, selected_lang)
            st.session_state['master_out'] = result

if 'master_out' in st.session_state:
    st.divider()
    st.subheader("👁️ AI PREVIEW")
    st.text_area("GENERATED CONTENT PREVIEW", st.session_state['master_out'], height=350)
    
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['master_out'], selected_lang)
    st.download_button(
        label="📥 DOWNLOAD WORD (.DOCX)", 
        data=doc_file, 
        file_name=f"LP_{u_topic.upper().replace(' ', '_')}_{selected_lang.upper()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- FOOTER SECTION ---
st.markdown("---") 
st.markdown(
    """
    <div style='text-align: center; color: grey; font-size: 0.8em;'>
        <p><b>SMART LESSON PLAN MASTER-PORTAL V1.1</b></p>
        <p>CONCEPTUALIZED BY: <b>HAJAH NURUL HAZIQAH @ HJH HARTINI HJ NORDIN</b></p>
        <p>© 2026 PTES Innovation | BSC H.M IN COMPUTER SCIENCE, UNIVERSITY OF STRATHCLYDE</p>
    </div>
    """,
    unsafe_allow_html=True
)
